import os
from docx import Document

# Initialize a Document
doc = Document()
doc.add_heading('Extracted File Contents', level=1)

# Define the file extensions you want to extract
extensions = {'.html', '.js', '.py'}

# Traverse the repository directory
for root, dirs, files in os.walk('.'):
    for file in files:
        if file == 'extract_files.py':
            continue
        ext = os.path.splitext(file)[1]
        if ext in extensions:
            file_path = os.path.join(root, file)
            doc.add_heading(file_path, level=2)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                doc.add_paragraph(content)

# Save the Document
doc.save('output.docx')


